import re
import logging
import requests
from io import BytesIO
from docx import Document
from docx.document import Document as DocumentObject
from typing import Optional, Dict, Any
from app.domain.repository.file_document_repository import IFileDocumentRepository

class WordManager(IFileDocumentRepository):
    def __init__(self, document_path: str, is_remote: Optional[bool] = False):
        self.document_path = document_path
        self.is_remote = is_remote
        self.logger = logging.getLogger(__name__)
        self.pattern = r"\{([A-Z0-9_]+)\}"
        
    def get_document(self) -> Document:
        """Obtiene el objeto Presentation desde una URL remota o una ruta local."""
        if self.is_remote:
            response = requests.get(self.document_path)
            response.raise_for_status()
            file_bytes = BytesIO(response.content)
            return Document(file_bytes)
        return Document(self.document_path)
    
    def analize_placeholders(self) -> Dict[str, Any]:

        place_holders = {}
        current_doc = self.get_document()

        for paragraph in current_doc.paragraphs:
            matches = re.findall(self.pattern, paragraph.text)
            for match in matches:
                place_holders[match] = {"type": "string"} 
        
        return place_holders

    def refill_document(self, fill_data: Dict[str, Any], output_path: Optional[str] = 'tmp/output.docx') -> None:
        current_doc = self.get_document()

        for paragraph in current_doc.paragraphs:
            for key, valor in fill_data.items():
                placeholder = f"{{{{{key}}}}}"
                if isinstance(valor, (str, int, float)):
                    paragraph.text = paragraph.text.replace(placeholder, str(valor))

        self.save_document(current_doc, output_path)

    def create_directories(self, file_path: str) -> None:
        from pathlib import Path
        path_object = Path(file_path)
        path_object.parent.mkdir(parents=True, exist_ok=True) 
        
    def save_document(self, document_content: DocumentObject, output_file_path: str) -> None:
        self.create_directories(output_file_path)
        document_content.save(output_file_path)

        self.logger.info(f"Documento guardado exitosamente en: {output_file_path}")

if __name__== '__main__':
    word_manager = WordManager('utility_test/TestDoc.docx')
    fill_data = word_manager.analize_placeholders()
    print("fill_data", fill_data)
    fill_data = {
        "PLACE": "San Juan de Lurigancho", 
        "SELECTED_DAY": "Miercoles", 
        "SELECTED_MONTH": "Febrero", 
        "SELECTED_YEAR": "2026", 
        "SELLER_ALL_NAMES": "Julio Romero Torres", 
        "BUYER_ALL_NAMES": "Ricardo Rojas", 
        "SERVICE_DESCRIPTION": "Desarrollo de una plataforma de agentes", 
        "PRICE_IN_LETTERS": "Mil quinientos", 
        "CURRENCY_SYMBOL": "$", 
        "TOTAL_MOUNT": "1500.00", 
        "CURRENCY": "DOLAR", 
        "METHOD_PAYS": "Transferencia, Yape, Plin"
    }
    word_manager.refill_document(fill_data)